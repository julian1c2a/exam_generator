import os
import random
from typing import List, Tuple, Any, Dict

# Solo importamos docx en la capa de "Presentación", no en la de Configuración
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

# ==============================================================================
# CAPA DE CONFIGURACIÓN (AGNÓSTICA)
# ==============================================================================

class ExamSpecs:
    """
    Contenedor centralizado de especificaciones de formato y contenido.
    Define el 'QUÉ' y el 'CÓMO' abstracto, sin depender de la tecnología de salida (Word, LaTeX, HTML).
    """

    # --- Configuración General del Documento ---
    TITLE = "Fundamentos de Electrónica"
    SUBTITLE = "Parte I: Electrónica Digital"
    STUDENT_FIELD = "Nombre: __________________________________________________  Fecha: ____________"

    FONT_MAIN = "Calibri"
    FONT_SIZE_MAIN = 11
    FONT_SIZE_TITLE = 14

    FONT_CODE = "Courier New"
    FONT_SIZE_CODE = 9

    # --- Ejercicio 1: Representación Numérica ---
    EX1_N_BITS = 8
    EX1_HEADERS = ['Id', 'Decimal', 'Binario Nat.', 'C2', 'Signo-Magnitud', 'BCD']
    EX1_ROW_LABELS = ['a', 'b', 'c', 'd']
    EX1_COL_WIDTH_CM = 2.5

    # --- Ejercicio 2: Karnaugh y Tablas ---
    EX2_TT_COL_WIDTH_CM = 1.0      # Ancho estrecho para columnas de Tabla de Verdad
    EX2_KMAP_CELL_DIM_CM = 1.2     # Dimensiones cuadradas (ancho y alto) para celdas del mapa
    EX2_KMAP_GRID_SIZE = 7         # Tamaño de la cuadrícula (7x7)
    EX2_TITLE_FONT_SIZE = 24       # Tamaño de la letra 'F' en la esquina

    # --- Ejercicio 3: Base de Datos de Escenarios ---
    EX3_SCENARIOS = [
        {
            "titulo": "Sistema de Seguridad de Bóveda Bancaria",
            "vars": ["A: Sensor Reloj (1=Laboral)", "B: Llave Director (1=Si)", "C: Llave Gerente (1=Si)", "D: Código (1=OK)"],
            "salida": "Z: Apertura",
            "logicas": [
                "La puerta se abre SI estamos en horario laboral Y se ha introducido el código correcto. ADEMÁS, por seguridad, fuera de horario laboral (A=0) la puerta también se puede abrir, pero SOLO si están presentes AMBAS llaves (Director y Gerente).",
                "La puerta se abre siempre que el código de seguridad sea correcto, SALVO que estemos fuera de horario laboral y falte alguna de las llaves (Director o Gerente).",
                "Para abrir la puerta se requiere estrictamente el código de seguridad. ADEMÁS, si estamos en horario laboral basta con eso, pero si NO es horario laboral, se requiere adicionalmente la llave del Director."
            ]
        },
        {
            "titulo": "Control de Reactor Químico Industrial",
            "vars": ["P: Presión (1=Alta)", "T: Temp (1=Alta)", "N: Nivel (1=Alto)", "M: Manual (1=ON)"],
            "salida": "E: Válvula Escape",
            "logicas": [
                "La válvula de escape debe abrirse SIEMPRE que se active el interruptor manual. AUTOMÁTICAMENTE, también debe abrirse si la Presión es alta Y (la Temperatura es alta O el Nivel es crítico).",
                "La válvula se abre si hay Presión alta, PERO solo si el Nivel también es crítico. Sin embargo, si la Temperatura es alta, la válvula se abre independientemente.",
                "Por seguridad, la válvula se abre si al menos DOS de los tres sensores (Presión, Temperatura, Nivel) están en estado alto. El interruptor Manual abre la válvula directamente."
            ]
        },
        {
            "titulo": "Sistema de Riego Inteligente",
            "vars": ["H: Humedad (1=Seco)", "L: Luz (1=Día)", "D: Depósito (1=Lleno)", "T: Temp (1=Calor)"],
            "salida": "R: Riego",
            "logicas": [
                "Riega SI (Seco Y Depósito lleno). ADEMÁS, si hace Calor excesivo, riega forzosamente siempre que haya agua.",
                "Nunca riega sin agua. Si hay agua, riega si está Seco, PERO evita regar de Día (L=1) salvo que haga Calor excesivo.",
                "Riega si está Seco. BLOQUEO: No riega si es de Día Y el depósito no está lleno. El Calor activa riego de emergencia."
            ]
        }
    ]

    # --- Ejercicio 5: Cronogramas ---
    EX5_CRONO_CYCLES = 6
    EX5_CRONO_HEIGHT_CM = 4.0

# ==============================================================================
# CAPA DE LÓGICA (Conversiones y Cálculos)
# ==============================================================================

def int_to_bin_str(val: int, bits: int) -> str:
    """Convierte un entero a string binario de N bits."""
    return format(val if val >= 0 else (1 << bits) + val, f'0{bits}b')

def int_to_bcd_str(val: int) -> str:
    """Convierte un entero (0-99) a su representación BCD de 8 bits."""
    if not (0 <= val <= 99): return "NR"
    tens = val // 10
    units = val % 10
    return f"{tens:04b} {units:04b}"

def int_to_sm_str(val: int, bits: int) -> str:
    """Convierte un entero a Signo-Magnitud."""
    if val >= 0:
        return format(val, f'0{bits}b')
    else:
        abs_val = abs(val)
        bin_abs = format(abs_val, f'0{bits - 1}b')
        return '1' + bin_abs

# ==============================================================================
# CAPA DE PRESENTACIÓN (Generación Word/Docx)
# ==============================================================================

def configurar_estilos(doc: Document):
    """Aplica las especificaciones de fuente de ExamSpecs al documento."""
    style = doc.styles['Normal']
    font = style.font
    font.name = ExamSpecs.FONT_MAIN
    font.size = Pt(ExamSpecs.FONT_SIZE_MAIN)

    try:
        style_code = doc.styles.add_style('CodeStyle', 1)
        style_code.font.name = ExamSpecs.FONT_CODE
        style_code.font.size = Pt(ExamSpecs.FONT_SIZE_CODE)
        style_code.paragraph_format.space_after = Pt(0)
        style_code.paragraph_format.space_before = Pt(0)
        style_code.paragraph_format.line_spacing = 1
        style_code.paragraph_format.keep_with_next = True
    except:
        pass

def agregar_ascii_art(doc: Document, art: str):
    p = doc.add_paragraph(art, style='CodeStyle')
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def agregar_encabezado(doc: Document):
    titulo = doc.add_heading(ExamSpecs.TITLE, 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitulo = doc.add_paragraph(ExamSpecs.SUBTITLE)
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo.runs[0].bold = True
    subtitulo.runs[0].font.size = Pt(ExamSpecs.FONT_SIZE_TITLE)

    doc.add_paragraph('_' * 85)
    doc.add_paragraph(ExamSpecs.STUDENT_FIELD)
    doc.add_paragraph('\n')

def generar_ejercicio_1(doc: Document) -> List[Tuple[str, int]]:
    doc.add_heading(f'Ejercicio 1: Sistemas de Representación (N = {ExamSpecs.EX1_N_BITS} bits)', level=1)

    p = doc.add_paragraph()
    p.add_run('a) Complete la tabla. ').bold = True
    p.add_run(f'El registro es de {ExamSpecs.EX1_N_BITS} bits. Si el número no es representable, escriba "NR".')

    table = doc.add_table(rows=len(ExamSpecs.EX1_ROW_LABELS) + 1, cols=len(ExamSpecs.EX1_HEADERS))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Cabeceras
    for i, text in enumerate(ExamSpecs.EX1_HEADERS):
        cell = table.rows[0].cells[i]
        cell.width = Cm(ExamSpecs.EX1_COL_WIDTH_CM)
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    valores_generados = []

    for i, label in enumerate(ExamSpecs.EX1_ROW_LABELS):
        row_cells = table.rows[i + 1].cells
        row_cells[0].text = f"{label})"
        row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        opciones_columna = [1, 1, 2, 3, 3, 4, 4, 5]
        col_idx = random.choice(opciones_columna)
        val = 0
        text_val = ""

        # Lógica de negocio (aleatoriedad)
        es_signed_col = col_idx in [1, 3, 4]
        signo = -1 if (es_signed_col and random.random() < 0.7) else 1
        n_bits = ExamSpecs.EX1_N_BITS

        if col_idx == 1: magnitud = random.randint(0, 200); val = signo * magnitud; text_val = str(val)
        elif col_idx == 2: val = random.randint(0, 255); text_val = int_to_bin_str(val, n_bits)
        elif col_idx == 3: limit = 128 if signo == -1 else 127; magnitud = random.randint(0, limit); val = signo * magnitud; text_val = int_to_bin_str(val, n_bits)
        elif col_idx == 4: magnitud = random.randint(0, 127); val = signo * magnitud; text_val = int_to_sm_str(val, n_bits)
        elif col_idx == 5: val = random.randint(0, 99); text_val = int_to_bcd_str(val)

        cell = row_cells[col_idx]
        cell.text = text_val
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        valores_generados.append((label, val))

    doc.add_paragraph('\n')
    return valores_generados

def generar_ejercicio_1_parte_b(doc: Document, valores: List[Tuple[str, int]]):
    if len(valores) < 2: return

    p_b = doc.add_paragraph()
    p_b.add_run('b) Realice las siguientes operaciones aritméticas ').bold = True
    p_b.add_run('utilizando los valores de la tabla anterior.')

    for i in range(1, 3):
        fila1 = random.choice(valores)
        fila2 = random.choice([x for x in valores if x != fila1])
        es_suma = random.choice([True, False])
        sistema = random.choice(['Binario Natural', 'Complemento a 2'])
        op_texto = "Suma" if es_suma else "Resta"
        signo = "+" if es_suma else "-"

        p_preg = doc.add_paragraph()
        p_preg.add_run(f'   {i}) {op_texto} en {sistema}: ').bold = True
        p_preg.add_run(f'Fila {fila1[0]} {signo} Fila {fila2[0]}')

        doc.add_paragraph('       Resultado Binario: __________________________')
        p_estado = doc.add_paragraph()
        p_estado.paragraph_format.left_indent = Cm(1.5)
        run = p_estado.add_run('¿Overflow? [   ]    ¿Underflow? [   ]    ¿Correcto? [   ]')
        run.italic = True
        doc.add_paragraph('')

def generar_ejercicio_2(doc: Document) -> Dict[str, Any]:
    doc.add_heading('Ejercicio 2: Diseño y Simplificación Lógica', level=1)

    # Lógica
    es_minterms = random.choice([True, False])
    tipo_canonico = "Minitérminos (Suma de Productos)" if es_minterms else "Maxitérminos (Producto de Sumas)"
    tipo_puerta = "NAND" if es_minterms else "NOR"
    target_val = 1 if es_minterms else 0
    default_val = 0 if es_minterms else 1
    outputs = [default_val] * 16

    for _ in range(random.randint(3, 6)):
        idx1 = random.randint(0, 15)
        idx2 = idx1 ^ (1 << random.randint(0, 3))
        outputs[idx1] = target_val
        outputs[idx2] = target_val

    doc.add_paragraph(f'Dada la función lógica F(A, B, C, D) definida por la siguiente tabla de verdad:')

    # --- TABLA DE VERDAD (Ancho controlado por Specs) ---
    table_tv = doc.add_table(rows=17, cols=5)
    table_tv.style = 'Table Grid'
    table_tv.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_tv.allow_autofit = False

    headers_tv = ['A', 'B', 'C', 'D', 'F']
    # Cabeceras
    for idx, text in enumerate(headers_tv):
        cell = table_tv.rows[0].cells[idx]
        cell.width = Cm(ExamSpecs.EX2_TT_COL_WIDTH_CM) # <-- USO DE SPEC
        cell.text = text
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].bold = True

    # Datos
    for i in range(16):
        row_cells = table_tv.rows[i+1].cells
        for cell in row_cells:
            cell.width = Cm(ExamSpecs.EX2_TT_COL_WIDTH_CM) # <-- USO DE SPEC

        bin_vals = f"{i:04b}"
        for j in range(4):
            row_cells[j].text = bin_vals[j]
            row_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[4].text = str(outputs[i])
        row_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[4].paragraphs[0].runs[0].bold = True

    doc.add_paragraph('\nUtilice el siguiente esquema para simplificar la función. (Nota: La numeración incluye opciones correctas e incorrectas, táchese la que no proceda).')

    # --- MAPA DE KARNAUGH (Dimensiones controladas por Specs) ---
    grid_size = ExamSpecs.EX2_KMAP_GRID_SIZE
    table_k = doc.add_table(rows=grid_size, cols=grid_size)
    table_k.style = 'Table Grid'
    table_k.alignment = WD_TABLE_ALIGNMENT.CENTER
    table_k.allow_autofit = False

    # Aplicar dimensiones cuadradas a toda la rejilla
    for r in range(grid_size):
        table_k.rows[r].height = Cm(ExamSpecs.EX2_KMAP_CELL_DIM_CM) # <-- USO DE SPEC
        for c in range(grid_size):
            table_k.cell(r, c).width = Cm(ExamSpecs.EX2_KMAP_CELL_DIM_CM) # <-- USO DE SPEC

    # Layout de la tabla (Trampa)
    cell_f = table_k.cell(0, 0)
    cell_f.merge(table_k.cell(2, 2))
    cell_f.text = "F"
    cell_f.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_f.paragraphs[0].runs[0].bold = True
    cell_f.paragraphs[0].runs[0].font.size = Pt(ExamSpecs.EX2_TITLE_FONT_SIZE)
    cell_f.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER

    cell_cd = table_k.cell(0, 3)
    cell_cd.merge(table_k.cell(0, 6))
    cell_cd.text = "CD ="
    cell_cd.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_cd.paragraphs[0].runs[0].bold = True

    cell_ab = table_k.cell(3, 0)
    cell_ab.merge(table_k.cell(6, 0))
    cell_ab.text = "AB ="
    cell_ab.vertical_alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_ab.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell_ab.paragraphs[0].runs[0].bold = True

    gray_code = ["00", "01", "11", "10"]
    bin_code = ["00", "01", "10", "11"]

    # Relleno de cabeceras (trampa vs correcto)
    for i in range(4):
        # Columnas
        table_k.cell(1, i+3).text = bin_code[i]
        table_k.cell(1, i+3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table_k.cell(2, i+3).text = gray_code[i]
        table_k.cell(2, i+3).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Filas
        table_k.cell(i+3, 1).text = bin_code[i]
        table_k.cell(i+3, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        table_k.cell(i+3, 2).text = gray_code[i]
        table_k.cell(i+3, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('\nSe pide:')
    doc.add_paragraph(f'a) Obtener la expresión canónica mediante {tipo_canonico}.', style='List Number')
    doc.add_paragraph('b) Identificar la numeración correcta del mapa (tachando la incorrecta) y simplificar.', style='List Number')
    doc.add_paragraph(f'c) Implementar el circuito simplificado utilizando EXCLUSIVAMENTE puertas {tipo_puerta}.', style='List Number')
    doc.add_paragraph('\n' * 2)
    return {"outputs": outputs}

def generar_ejercicio_3(doc: Document) -> Dict[str, Any]:
    doc.add_heading('Ejercicio 3: Problema de Diseño Lógico', level=1)

    # Uso de la Base de Datos de Specs
    escenario = random.choice(ExamSpecs.EX3_SCENARIOS)
    logica_texto = random.choice(escenario["logicas"])

    doc.add_paragraph(f"Contexto: {escenario['titulo']}").bold = True
    for var in escenario["vars"]: doc.add_paragraph(var, style='List Bullet')
    doc.add_paragraph(f"Lógica: {logica_texto}").italic = True

    doc.add_paragraph('\nSe pide: Tabla de Verdad, Karnaugh y Esquema Lógico.')
    doc.add_paragraph('\n' * 3)
    return {"escenario": escenario}

def generar_ejercicio_4(doc: Document) -> Dict[str, Any]:
    doc.add_heading('Ejercicio 4: Análisis de Bloques MSI', level=1)

    tipo_bloque = random.choice(['MUX', 'COMPARADOR', 'SUMADOR'])
    datos_retorno = {"tipo": tipo_bloque}

    doc.add_paragraph(f'Dado el siguiente componente: {tipo_bloque}').bold = True

    if tipo_bloque == 'MUX':
        inputs_mux = [random.randint(0, 1) for _ in range(16)]
        inputs_str = [str(x) for x in inputs_mux]
        col1 = inputs_str[:8]; col2 = inputs_str[8:]

        ascii_art = (
            f"          +-----------------------+\n"
            f" I0 ={col1[0]} ---|                       |--- Y\n"
            f" I1 ={col1[1]} ---|                       |--- Y_neg\n"
            f" ...      |       MUX 16:1        |\n"
            f" I7 ={col1[7]} ---|                       |\n"
            f"          |                       |\n"
            f" I8 ={col2[0]} ---|                       |\n"
            f" ...      |                       |\n"
            f" I15={col2[7]} ---|                       |\n"
            f"          +-----------------------+\n"
            f"             |   |   |   |     |\n"
            f"             S3  S2  S1  S0    E\n"
        )
        agregar_ascii_art(doc, ascii_art)
        doc.add_paragraph(f'Entradas I0..I15: {inputs_mux}')

        doc.add_paragraph('\nDetermine salidas Y y Y_neg:')
        for i in range(1, 4):
            addr_val = random.randint(0, 15)
            enable = random.choice([0, 1])
            enable_str = "0 (Activo)" if enable == 0 else "1 (Inactivo)"
            doc.add_paragraph(f'{i}) E={enable_str}, Dir={addr_val:04b}.', style='List Number')

    elif tipo_bloque == 'COMPARADOR':
        val_a = random.randint(0, 15); val_b = random.randint(0, 15)
        cascada = [random.choice([0, 1]) for _ in range(3)]

        ascii_art = (
            f"             +-------------------+\n"
            f"  A = {val_a:04b}  --|                   |--- A > B\n"
            f"             |    COMPARADOR     |\n"
            f"  B = {val_b:04b}  --|      4 BITS       |--- A = B\n"
            f"             |                   |\n"
            f"  I(> )={cascada[0]} --|                   |--- A < B\n"
            f"  I(= )={cascada[1]} --|                   |\n"
            f"  I(< )={cascada[2]} --|                   |\n"
            f"             +-------------------+\n"
        )
        agregar_ascii_art(doc, ascii_art)
        doc.add_paragraph('Determine > = <.')

    elif tipo_bloque == 'SUMADOR':
        val_a = random.randint(0, 15); val_b = random.randint(0, 15); c_in = random.choice([0, 1])
        ascii_art = (
            f"             +-------------------+\n"
            f"  A = {val_a:04b}  --|                   |--- S (Suma)\n"
            f"             |      SUMADOR      |\n"
            f"  B = {val_b:04b}  --|      4 BITS       |--- Cout\n"
            f"             |                   |\n"
            f"  Cin = {c_in}  ---|                   |\n"
            f"             +-------------------+\n"
        )
        agregar_ascii_art(doc, ascii_art)
        doc.add_paragraph('Determine S, Cout y Overflow.')

    doc.add_paragraph('\n' * 2)
    return datos_retorno

def generar_ejercicio_5(doc: Document) -> Dict[str, Any]:
    doc.add_heading('Ejercicio 5: Análisis de Sistemas Secuenciales', level=1)

    ff_type = random.choice(['JK', 'D', 'T'])
    edge = random.choice(['Subida', 'Bajada'])
    has_async = random.choice([True, False])

    async_text = "Sin Async"
    if has_async:
        atype = random.choice(['Preset', 'Clear'])
        alevel = random.choice(['1', '0'])
        async_text = f"Async {atype} activa a {alevel}"

    logic_type = 'SHIFT' if ff_type == 'D' else 'COUNTER'

    doc.add_paragraph(f"Sistema síncrono por flanco {edge}. FF: {ff_type}. {async_text}. Lógica: {logic_type}")

    # Esquema simple
    clock_sym = "> " if edge == 'Subida' else "o>"
    ascii_ckt = (
        f"      +-------+          +-------+\n"
        f" E ---|{ff_type:<5}  Q|---Q0----|{ff_type:<5}  Q|--- Q1\n"
        f"CLK --|{clock_sym}     |          |{clock_sym}     |\n"
        f"      +-------+          +-------+\n"
    )
    agregar_ascii_art(doc, ascii_ckt)

    # Cronograma
    cycles = ExamSpecs.EX5_CRONO_CYCLES
    clk_wave =  "__|--|__|--|__|--|__|--|__|--|__|--"
    input_bits = [random.randint(0,1) for _ in range(cycles*2)]
    input_wave = "".join(["~~~" if b else "___" for b in input_bits])

    crono_art = (
        f"       t0 t1 t2 t3 t4 t5 t6 ...\n"
        f" CLK:  {clk_wave}\n"
        f"  E :  {input_wave}\n"
        f" Q0 :  ...................................\n"
        f" Q1 :  ...................................\n"
    )

    agregar_ascii_art(doc, crono_art)

    doc.add_paragraph('a) Ecuaciones de entrada. b) Completar cronograma.')

    return {"ff": ff_type}

# ==============================================================================
# MAIN
# ==============================================================================

def guardar_documento(doc: Document):
    nombre_final = 'Examen_Electronica_Digital.docx'
    try:
        doc.save(nombre_final)
        print(f"\n✅ ¡ÉXITO! Examen generado correctamente.")
        print(f"📂 Archivo: {os.path.abspath(nombre_final)}")
    except PermissionError:
        nuevo_nombre = f'Examen_Electronica_Digital_{random.randint(100, 999)}.docx'
        doc.save(nuevo_nombre)
        print(f"✅ ¡GUARDADO! Copia nueva: {nuevo_nombre}")

def main():
    doc = Document()
    configurar_estilos(doc)
    agregar_encabezado(doc)

    valores_ej1 = generar_ejercicio_1(doc)
    generar_ejercicio_1_parte_b(doc, valores_ej1)

    valores_ej2 = generar_ejercicio_2(doc)
    valores_ej3 = generar_ejercicio_3(doc)
    valores_ej4 = generar_ejercicio_4(doc)
    valores_ej5 = generar_ejercicio_5(doc)

    guardar_documento(doc)

if __name__ == "__main__":
    main()