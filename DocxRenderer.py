import os
import random
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from exam_model import *

class DocxRenderer:
    def __init__(self):
        self.doc = Document()
        self._configurar_estilos()

    def _configurar_estilos(self):
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)

    def render_examen(self, ej1: Exercise1Data, ej2: Exercise2Data, ej3: Exercise3Data, ej4: Exercise4Data, ej5: Exercise5Data, filename: str):
        self._agregar_encabezado()
        self._render_ej1(ej1)
        self._render_ej2(ej2)
        self._render_ej3(ej3)
        self._render_ej4(ej4)
        self._render_ej5(ej5)
        self._guardar_documento(filename)

    def _agregar_encabezado(self):
        titulo = self.doc.add_heading('Fundamentos de Electrónica', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitulo = self.doc.add_paragraph('Parte I: Electrónica Digital')
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitulo.runs[0].bold = True
        subtitulo.runs[0].font.size = Pt(14)

        self.doc.add_paragraph('_' * 85)
        self.doc.add_paragraph('Nombre: __________________________________________________  Fecha: ____________')
        self.doc.add_paragraph('\n')

    def _render_ej1(self, data: Exercise1Data):
        self.doc.add_heading(f'Ejercicio 1: Sistemas de Representación (N = {data.n_bits} bits)', level=1)

        p = self.doc.add_paragraph()
        p.add_run('a) Complete la tabla. ').bold = True
        p.add_run(f'El registro es de {data.n_bits} bits. Si el número no es representable, escriba "NR".')

        headers = ['Id', 'Decimal', 'Binario Nat.', 'C2', 'Signo-Magnitud', 'BCD']
        table = self.doc.add_table(rows=len(data.rows) + 1, cols=6)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Cabecera
        hdr_cells = table.rows[0].cells
        for i, text in enumerate(headers):
            run = hdr_cells[i].paragraphs[0].add_run(text)
            run.bold = True
            run.font.size = Pt(9)
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[i].width = Cm(2.5)

        # Filas
        for i, row in enumerate(data.rows):
            row_cells = table.rows[i + 1].cells
            row_cells[0].text = f"{row.label})"
            row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Rellenar solo la celda objetivo
            cell = row_cells[row.target_col_idx]
            cell.text = row.target_val_str
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph('\n')

        # Parte B
        if data.operations:
            p_b = self.doc.add_paragraph()
            p_b.add_run('b) Realice las siguientes operaciones aritméticas:').bold = True
            
            for i, op in enumerate(data.operations, 1):
                self.doc.add_paragraph(f'   {i}) {op.op_type} en {op.system}:  Fila {op.operand1} {op.operator_symbol} Fila {op.operand2}')
            self.doc.add_paragraph('\n' * 1)

    def _render_ej2(self, data: Exercise2Data):
        self.doc.add_heading('Ejercicio 2: Diseño y Simplificación Lógica', level=1)
        self.doc.add_paragraph('Dada la función lógica definida por la siguiente tabla de verdad (ver adjunto):')

        # Placeholder para tabla de verdad (en Word es difícil generarla dinámica bonita sin mucho código,
        # así que dejamos el espacio o una tabla genérica)
        cuadro = self.doc.add_table(rows=1, cols=1)
        cuadro.style = 'Table Grid'
        cuadro.rows[0].cells[0].height = Cm(3)
        p = cuadro.rows[0].cells[0].paragraphs[0]
        p.add_run('\n[PEGAR TABLA DE VERDAD AQUÍ]\n').italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph('\nSe pide:')
        self.doc.add_paragraph(f'a) Obtener la expresión canónica ({data.canon_type}).', style='List Number')
        self.doc.add_paragraph('b) Simplificar mediante Mapas de Karnaugh.', style='List Number')
        self.doc.add_paragraph(f'c) Implementar con puertas universales ({data.gate_type}) según simplificación.', style='List Number')
        self.doc.add_paragraph('\n' * 2)

    def _render_ej3(self, data: Exercise3Data):
        self.doc.add_heading('Ejercicio 3: Resolución de Problema de Sistema', level=1)
        p = self.doc.add_paragraph()
        p.add_run(f'Contexto: {data.title}. ').bold = True
        self.doc.add_paragraph(f'Lógica: {data.logic_description}')
        
        self.doc.add_paragraph('\nVariables:', style='No Spacing')
        for v in data.variables:
            self.doc.add_paragraph(f'- {v}', style='No Spacing')
        self.doc.add_paragraph(f'- Salida: {data.output_name}', style='No Spacing')

        self.doc.add_paragraph('\nPasos requeridos:')
        self.doc.add_paragraph('1. Tabla de Verdad.', style='List Bullet')
        self.doc.add_paragraph('2. Simplificación (Karnaugh).', style='List Bullet')
        self.doc.add_paragraph('3. Esquema lógico.', style='List Bullet')
        self.doc.add_paragraph('\n' * 2)

    def _render_ej4(self, data: Exercise4Data):
        self.doc.add_heading('Ejercicio 4: Análisis de Bloques', level=1)
        self.doc.add_paragraph(f'Dado el siguiente esquema lógico ({data.block_type}):')

        cuadro = self.doc.add_table(rows=1, cols=1)
        cuadro.style = 'Table Grid'
        cuadro.rows[0].cells[0].height = Cm(4)
        p = cuadro.rows[0].cells[0].paragraphs[0]
        p.add_run(f'\n[INSERTAR IMAGEN DE {data.block_type} AQUÍ]\n').italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph('\nResponda:')
        self.doc.add_paragraph('a) Determine las salidas para las entradas indicadas.', style='List Number')
        self.doc.add_paragraph('b) Analice la función lógica global.', style='List Number')
        self.doc.add_paragraph('\n' * 2)

    def _render_ej5(self, data: Exercise5Data):
        self.doc.add_heading('Ejercicio 5: Sistemas Secuenciales', level=1)
        
        desc = f"FF {data.ff_type}, disparo por {data.edge_type}. "
        if data.has_async:
            desc += f"Entrada asíncrona {data.async_type} activa a nivel {data.async_level}."
        
        self.doc.add_paragraph(f'Analice el cronograma adjunto para el sistema secuencial dado ({desc}).')

        cuadro = self.doc.add_table(rows=1, cols=1)
        cuadro.style = 'Table Grid'
        cuadro.rows[0].cells[0].height = Cm(4)
        p = cuadro.rows[0].cells[0].paragraphs[0]
        p.add_run('\n[INSERTAR CRONOGRAMA AQUÍ]\n').italic = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_paragraph('\nSe pide:')
        self.doc.add_paragraph('a) Completar el cronograma (salidas Q0, Q1).', style='List Number')
        self.doc.add_paragraph('b) Determinar la secuencia de estados.', style='List Number')

    def _guardar_documento(self, filename: str):
        try:
            self.doc.save(filename)
            print(f"✅ Documento Word guardado: {os.path.abspath(filename)}")
        except PermissionError:
            print(f"⚠️ Error: No se pudo guardar '{filename}'. Probablemente esté abierto.")
            nuevo_nombre = f"Examen_Copia_{random.randint(100,999)}.docx"
            self.doc.save(nuevo_nombre)
            print(f"✅ Guardado como copia: {os.path.abspath(nuevo_nombre)}")
